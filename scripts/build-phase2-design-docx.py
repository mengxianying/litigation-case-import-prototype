from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / 'docs/superpowers/specs/2026-07-19-case-import-phase2-design.md'
OUTPUT = ROOT / 'docs/superpowers/specs/案件导入二期设计说明.docx'

INK = '243142'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
LIGHT_BLUE = 'EAF3FB'
LIGHT_GRAY = 'F2F4F7'
BORDER = 'D9E2EC'
MUTED = '5F6B7A'
FONT = 'PingFang SC'


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{margin}'))
        if node is None:
            node = OxmlElement(f'w:{margin}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_twips))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = table_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        table_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    tbl_w = table_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        table_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    tr_pr.append(header)


def add_bottom_border(paragraph, color=BLUE, size='14'):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '7')
    bottom.set(qn('w:color'), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def apply_runs(paragraph, text, bold=False, color=None, size=None):
    token_pattern = re.compile(r'(\*\*.*?\*\*|`.*?`)')
    pieces = token_pattern.split(text)
    for piece in pieces:
        if not piece:
            continue
        run = paragraph.add_run()
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        run.bold = bold
        if piece.startswith('**') and piece.endswith('**'):
            run.text = piece[2:-2]
            run.bold = True
        elif piece.startswith('`') and piece.endswith('`'):
            run.text = piece[1:-1]
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.font.color.rgb = RGBColor(177, 76, 50)
            run.font.size = Pt(9)
        else:
            run.text = piece
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if size:
            run.font.size = Pt(size)


def clean_cell_text(text):
    return text.strip().replace('\\n', '\n')


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith('|'):
        row = [cell.strip() for cell in lines[index].strip().strip('|').split('|')]
        if not all(re.fullmatch(r':?-{3,}:?', cell.replace(' ', '')) for cell in row):
            rows.append(row)
        index += 1
    return rows, index


def add_table(doc, rows):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [''] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=column_count)
    table.style = 'Table Grid'
    usable = 9360
    distributions = {
        1: [9360],
        2: [2700, 6660],
        3: [1900, 2750, 4710],
        4: [1700, 2350, 3180, 2130],
        5: [1320, 1800, 2140, 1800, 2300],
    }
    widths = distributions.get(column_count, [usable // column_count] * column_count)
    if sum(widths) != usable:
        widths[-1] += usable - sum(widths)
    set_table_geometry(table, widths)
    for row_index, values in enumerate(normalized):
        for cell, value in zip(table.rows[row_index].cells, values):
            cell.text = ''
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            apply_runs(paragraph, clean_cell_text(value), bold=row_index == 0, color=DARK_BLUE if row_index == 0 else INK, size=9.5)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, 'FAFCFE')
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_flow(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    apply_runs(p, '一期历史批次补材流程', bold=True, color=DARK_BLUE, size=11)
    steps = [
        '待初始化（一期）\n案件已生效\n导入完成',
        '补充历史材料\n上传材料包',
        '材料处理中\n固化规则快照',
        '材料齐全\n处理详情查看解析',
        '待补必填 / 待补非必填\n处理详情批量补传',
        '案件继续已生效\n导入完成不变',
    ]
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    set_table_geometry(table, [3120, 3120, 3120])
    for index, text in enumerate(steps):
        cell = table.rows[index // 3].cells[index % 3]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for line_index, line in enumerate(text.split('\n')):
            run = p.add_run(line)
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            run.bold = line_index == 0
            if line_index < len(text.split('\n')) - 1:
                run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    styles = {
        'Heading 1': (16, BLUE, 16, 8),
        'Heading 2': (13, BLUE, 12, 6),
        'Heading 3': (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in styles.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = '案件导入二期设计说明 | V2.0'
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = FONT
    header.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('第 ')
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(footer)
    run = footer.add_run(' 页')
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build():
    doc = Document()
    configure_document(doc)
    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    index = 0
    title_written = False
    in_mermaid = False
    mermaid_lines = []
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if line.startswith('```'):
            if line == '```mermaid':
                in_mermaid = True
                mermaid_lines = []
            elif in_mermaid:
                add_flow(doc)
                in_mermaid = False
            index += 1
            continue
        if in_mermaid:
            mermaid_lines.append(raw)
            index += 1
            continue
        if not line or line == '---':
            index += 1
            continue
        if line.startswith('|'):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            apply_runs(p, line[2:].strip(), bold=True, color=DARK_BLUE, size=23)
            add_bottom_border(p)
            title_written = True
        elif line.startswith('## '):
            doc.add_paragraph(line[3:].strip(), style='Heading 1')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:].strip(), style='Heading 2')
        elif line.startswith('#### '):
            doc.add_paragraph(line[5:].strip(), style='Heading 3')
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            apply_runs(p, re.sub(r'^\d+\. ', '', line), color=INK, size=10.5)
        elif line.startswith('- [ ] '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            apply_runs(p, '□ ' + line[6:], color=INK, size=10.5)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            apply_runs(p, line[2:], color=INK, size=10.5)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            apply_runs(p, line, color=INK, size=10.5)
        index += 1
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    build()
